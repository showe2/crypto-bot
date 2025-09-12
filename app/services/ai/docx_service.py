import tempfile
import os
from typing import Dict, Any, Optional
from datetime import datetime
from loguru import logger

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Install with: pip install python-docx")


class DocxReportService:
    """DOCX report generator for token analysis"""

    async def generate_run_docx(self, run_data: Dict[str, Any], profile_type: str) -> Optional[bytes]:
        """Generate DOCX report from run data based on profile type"""
        
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        try:
            run_id = run_data.get("run_id", "unknown")
            logger.info(f"📄 Generating {profile_type} DOCX for run: {run_id}")
            
            # Create fresh document
            doc = Document()
            
            # Add run info section (common to all types)
            self._add_run_info_section(doc, run_data, profile_type)
            
            # Route to specific generator based on profile type
            if profile_type == "pump":
                self._build_pump_report(doc, run_data)
            elif profile_type == "discovery":
                self._build_discovery_report(doc, run_data)
            elif profile_type in ["whale", "twitter", "listing"]:
                self._build_generic_profile_report(doc, run_data, profile_type)
            else:
                raise ValueError(f"Unsupported profile type: {profile_type}")
            
            # Convert to bytes
            temp_path = None
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    temp_path = tmp.name
                    doc.save(temp_path)
                
                with open(temp_path, 'rb') as f:
                    docx_bytes = f.read()
                    logger.info(f"✅ {profile_type} DOCX generated ({len(docx_bytes)} bytes)")
                    return docx_bytes
                    
            finally:
                if temp_path and os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            logger.error(f"❌ {profile_type} DOCX generation failed: {str(e)}")
            raise

    def _add_run_info_section(self, doc: Document, run_data: Dict[str, Any], profile_type: str):
        """Add common run information section to all DOCX types"""
        try:
            # Title
            heading = doc.add_heading(f'{profile_type.title()} Analysis Report', 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Run information table
            doc.add_heading('Run Information', level=1)
            
            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'
            
            # Header
            header_row = info_table.rows[0]
            header_row.cells[0].text = "Property"
            header_row.cells[1].text = "Value"
            
            # Run details
            run_info = [
                ("Run ID", run_data.get("run_id", "Unknown")),
                ("Analysis Type", profile_type.title()),
                ("Timestamp", self._format_timestamp(run_data.get("timestamp"))),
                ("Processing Time", f"{run_data.get('processing_time', 0):.2f} seconds"),
                ("Status", run_data.get("status", "Unknown").title()),
                ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC"))
            ]
            
            # Add filters for pump type
            if profile_type == "pump" and run_data.get("filters"):
                filters = run_data["filters"]
                run_info.extend([
                    ("Liquidity Range", f"${filters.get('liqMin', 0):,} - ${filters.get('liqMax', 0):,}"),
                    ("Market Cap Range", f"${filters.get('mcapMin', 0):,} - ${filters.get('mcapMax', 0):,}"),
                    ("Volume Range", f"${filters.get('volMin', 0):,} - ${filters.get('volMax', 0):,}"),
                    ("Time Range", f"{filters.get('timeMin', 0)} - {filters.get('timeMax', 0)} minutes")
                ])
            
            # Add discovery token info
            if profile_type == "discovery" and run_data.get("results"):
                results = run_data["results"]
                if len(results) > 0:
                    result = results[0]
                    run_info.extend([
                        ("Token Address", result.get("token_address", "Unknown")),
                        ("Token Symbol", result.get("token_symbol", "Unknown")),
                        ("Token Name", result.get("token_name", "Unknown"))
                    ])
            
            # Add rows to table
            for prop, value in run_info:
                row = info_table.add_row()
                row.cells[0].text = str(prop)
                row.cells[1].text = str(value)
            
            logger.info(f"Added run info section with {len(run_info)} properties")
            
        except Exception as e:
            logger.warning(f"Failed to add run info section: {e}")
            doc.add_paragraph("Run information section unavailable")

    def _format_timestamp(self, timestamp) -> str:
        """Format timestamp for display"""
        try:
            if not timestamp:
                return "Unknown"
            
            # Handle both unix timestamp and datetime string
            if isinstance(timestamp, (int, float)):
                dt = datetime.fromtimestamp(timestamp)
            else:
                dt = datetime.fromisoformat(str(timestamp).replace('Z', '+00:00'))
                if dt.tzinfo:
                    dt = dt.replace(tzinfo=None)
            
            return dt.strftime("%Y-%m-%d %H:%M:%S UTC")
            
        except Exception:
            return str(timestamp) if timestamp else "Unknown"
        
    def _build_pump_report(self, doc: Document, run_data: Dict[str, Any]):
        """Build pump analysis report with candidates table"""
        try:
            # Extract pump results
            results = run_data.get("results", [])
            candidates_found = run_data.get("candidates_found", len(results))
            snapshots_analyzed = run_data.get("snapshots_analyzed", 0)
            
            # Extract filters
            filters = run_data.get("filters", {})
            
            # Summary section
            doc.add_heading('Pump Analysis Summary', level=1)
            
            summary_p = doc.add_paragraph()
            summary_p.add_run('Snapshots Analyzed: ').bold = True
            summary_p.add_run(f'{snapshots_analyzed:,}\n')
            summary_p.add_run('Candidates Found: ').bold = True
            summary_p.add_run(f'{candidates_found}\n')
            summary_p.add_run('Top Candidates Shown: ').bold = True
            summary_p.add_run(f'{min(len(results), 5)}\n')
            
            # Add filter summary
            if filters:
                doc.add_heading('Filter Criteria Applied', level=2)
                filter_p = doc.add_paragraph()
                filter_p.add_run('Liquidity Range: ').bold = True
                filter_p.add_run(f"${filters.get('liqMin', 0):,.0f} - ${filters.get('liqMax', 0):,.0f}\n")
                filter_p.add_run('Market Cap Range: ').bold = True
                filter_p.add_run(f"${filters.get('mcapMin', 0):,.0f} - ${filters.get('mcapMax', 0):,.0f}\n")
                filter_p.add_run('Time Range: ').bold = True
                filter_p.add_run(f"{filters.get('timeMin', 0)} - {filters.get('timeMax', 0)} minutes\n")
            
            if not results:
                doc.add_paragraph("No pump candidates found matching the specified criteria.")
                return
            
            # Candidates table
            doc.add_heading('Top Pump Candidates', level=1)
            
            # Create table with headers - Updated for your data structure
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            
            # Headers
            headers = ["Rank", "Token Name", "Contract", "Liquidity", "Vol 5m", "Vol 60m", 
                    "Market Cap", "Whales 1h", "Pump Score", "AI Analysis"]
            
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                header_row.cells[i].paragraphs[0].runs[0].bold = True
            
            # Add candidate rows (top 5)
            for candidate in results[:5]:
                row = table.add_row()
                
                # Extract candidate data safely
                rank = candidate.get("rank", "N/A")
                name = candidate.get("name", "Unknown")
                contract = candidate.get("mint", "Unknown")
                liq = candidate.get("liq", 0)
                vol5 = candidate.get("vol5", 0)
                vol60 = candidate.get("vol60", 0)
                mcap = candidate.get("mcap", 0)
                pump_score = candidate.get("pump_score", 0)
                
                # Whale activity - handle your specific format
                whales_data = candidate.get("whales1h", {})
                if isinstance(whales_data, dict):
                    whale_count = whales_data.get("count", 0)
                    whale_inflow = whales_data.get("total_inflow_usd", 0)
                    whales_text = f"{whale_count} (+${whale_inflow:,.0f})" if whale_inflow > 0 else f"{whale_count}"
                else:
                    whales_text = str(whales_data)
                
                # AI analysis - decode Unicode if needed
                ai_analysis = candidate.get("ai", "No analysis")
                if isinstance(ai_analysis, str) and "\\u" in ai_analysis:
                    try:
                        ai_analysis = ai_analysis.encode().decode('unicode_escape')
                    except:
                        pass
                
                # Populate row
                row.cells[0].text = str(rank)
                row.cells[1].text = str(name)
                row.cells[2].text = str(contract)
                row.cells[3].text = f"${liq:,}" if isinstance(liq, (int, float)) else str(liq)
                row.cells[4].text = f"${vol5:,}" if isinstance(vol5, (int, float)) else str(vol5)
                row.cells[5].text = f"${vol60:,}" if isinstance(vol60, (int, float)) else str(vol60)
                row.cells[6].text = f"${mcap:,}" if isinstance(mcap, (int, float)) else str(mcap)
                row.cells[7].text = whales_text
                row.cells[8].text = f"{pump_score:.2f}" if isinstance(pump_score, (int, float)) else str(pump_score)
                row.cells[9].text = ai_analysis[:150] + "..." if len(ai_analysis) > 150 else ai_analysis
            
            # Add methodology section
            doc.add_heading('Analysis Methodology', level=1)
            
            method_text = """This pump analysis uses real-time snapshot data to identify tokens with high pump potential. 

    Key Metrics:
    - Pump Score: Calculated as (Volume × Volume/Liquidity Ratio) / Age in Minutes
    - Whale Activity: Number of large holders and their recent inflow activity
    - Liquidity Health: Token's available liquidity for trading
    - Volume Momentum: Recent 5-minute vs 60-minute volume comparison

    Filtering Criteria:
    - Market cap, liquidity, and volume ranges as specified
    - Token age and whale activity thresholds
    - Minimum pump score requirements

    Risk Warning: This analysis is for informational purposes only. Cryptocurrency trading involves significant risk of loss."""
            
            doc.add_paragraph(method_text)
            
            logger.info(f"Built pump report with {len(results[:5])} candidates")
            
        except Exception as e:
            logger.error(f"Error building pump report: {e}")
            doc.add_paragraph(f"Error generating pump report: {str(e)}")

    def _build_generic_profile_report(self, doc: Document, run_data: Dict[str, Any], profile_type: str):
        """Build generic profile report for whale, twitter, listing types"""
        try:
            doc.add_heading(f'{profile_type.title()} Analysis Results', level=1)
            
            results = run_data.get("results", [])
            
            if not results:
                doc.add_paragraph(f"No {profile_type} analysis results found.")
                return
            
            # Simple results table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Headers
            headers = ["Token", "Address", "Status", "Summary"]
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                header_row.cells[i].text = header
                header_row.cells[i].paragraphs[0].runs[0].bold = True
            
            # Add result rows
            for result in results[:10]:  # Show top 10
                row = table.add_row()
                row.cells[0].text = result.get("token_symbol", "Unknown")
                row.cells[1].text = result.get("token_address", "Unknown")[:16] + "..."
                row.cells[2].text = "Success" if result.get("success") else "Failed"
                row.cells[3].text = f"{profile_type.title()} analysis completed"
            
            logger.info(f"Built {profile_type} report with {len(results[:10])} results")
            
        except Exception as e:
            logger.error(f"Error building {profile_type} report: {e}")
            doc.add_paragraph(f"Error generating {profile_type} report: {str(e)}")

    def _build_discovery_report(self, doc: Document, run_data: Dict[str, Any]):
        """Build discovery analysis report using existing comprehensive format"""
        try:
            # Extract discovery analysis data from run results
            results = run_data.get("results", [])
            
            if not results:
                doc.add_paragraph("No discovery analysis results found.")
                return
            
            # Get the analysis result from first result
            first_result = results[0]
            analysis_data = first_result.get("analysis_result")
            
            if not analysis_data:
                doc.add_paragraph("Discovery analysis data not available.")
                return
            
            # Extract token info
            token_symbol = first_result.get("token_symbol", "Unknown")
            token_name = first_result.get("token_name", "Unknown Token") 
            token_address = first_result.get("token_address", "N/A")
            
            logger.info(f"Building discovery report for {token_symbol} ({token_name})")
            
            # Token Info Section
            doc.add_heading('Token Information', level=1)
            
            token_p = doc.add_paragraph()
            token_p.add_run('Token Name: ').bold = True
            token_p.add_run(f'{token_name}\n')
            token_p.add_run('Symbol: ').bold = True
            token_p.add_run(f'{token_symbol}\n')
            token_p.add_run('Contract Address: ').bold = True
            token_p.add_run(f'{token_address}\n')
            
            # Use existing methods but extract data from analysis_data
            self._add_discovery_ai_verdict_section(doc, analysis_data)
            self._add_discovery_scoring_table(doc, analysis_data)
            self._add_discovery_pros_cons_section(doc, analysis_data)
            self._add_discovery_ai_reasoning_section(doc, analysis_data)
            self._add_discovery_market_data_section(doc, analysis_data)
            self._add_discovery_security_section(doc, analysis_data)
            
            logger.info(f"Discovery report built successfully for {token_symbol}")
            
        except Exception as e:
            logger.error(f"Error building discovery report: {e}")
            doc.add_paragraph(f"Error generating discovery report: {str(e)}")

    def _add_discovery_ai_verdict_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add AI verdict section for discovery report"""
        try:
            doc.add_heading('AI Investment Verdict', level=1)
            
            overall_analysis = analysis_data.get("overall_analysis", {})
            ai_analysis = analysis_data.get("ai_analysis", {})
            
            # Get verdict and score
            verdict_decision = "UNKNOWN"
            if ai_analysis and ai_analysis.get("recommendation"):
                ai_rec = ai_analysis["recommendation"]
                verdict_mapping = {"BUY": "GO", "CONSIDER": "WATCH", "HOLD": "WATCH", "CAUTION": "WATCH", "AVOID": "NO"}
                verdict_decision = verdict_mapping.get(ai_rec, "WATCH")
            else:
                score = overall_analysis.get("score", 0)
                if score >= 80:
                    verdict_decision = "GO"
                elif score >= 60:
                    verdict_decision = "WATCH" 
                else:
                    verdict_decision = "NO"
            
            # Display verdict
            p = doc.add_paragraph()
            p.add_run('INVESTMENT VERDICT: ').bold = True
            
            verdict_run = p.add_run(verdict_decision)
            verdict_run.bold = True
            
            # Color coding
            if verdict_decision == "GO":
                verdict_run.font.color.rgb = RGBColor(34, 197, 94)  # Green
            elif verdict_decision == "WATCH":
                verdict_run.font.color.rgb = RGBColor(245, 158, 11)  # Yellow
            else:  # NO
                verdict_run.font.color.rgb = RGBColor(239, 68, 68)  # Red
            
            # Add AI score
            ai_score = ai_analysis.get("ai_score", 0) if ai_analysis else 0
            overall_score = overall_analysis.get("score", 0)
            
            p.add_run(f'\nAI Score: ').bold = True
            p.add_run(f'{ai_score}/100')
            
            p.add_run(f'\nOverall Score: ').bold = True
            p.add_run(f'{overall_score}/100')
            
        except Exception as e:
            logger.warning(f"Failed to add discovery AI verdict: {e}")
            doc.add_paragraph("AI verdict section unavailable")

    def _add_discovery_scoring_table(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add scoring table for discovery report"""
        try:
            doc.add_heading('Risk Metrics Analysis', level=1)
            
            ai_analysis = analysis_data.get("ai_analysis", {})
            service_responses = analysis_data.get("service_responses", {})
            
            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Headers
            header_row = table.rows[0]
            header_row.cells[0].text = "Risk Metric"
            header_row.cells[1].text = "Value" 
            header_row.cells[2].text = "AI Assessment"
            
            # Extract key metrics
            metrics_data = []
            
            # Market data from Birdeye
            birdeye_price = service_responses.get("birdeye", {}).get("price", {})
            if birdeye_price:
                market_cap = birdeye_price.get("market_cap")
                if market_cap:
                    metrics_data.append(("Market Cap", f"${float(market_cap):,.0f}", "Analyzed"))
                
                liquidity = birdeye_price.get("liquidity")  
                if liquidity:
                    metrics_data.append(("Liquidity", f"${float(liquidity):,.0f}", "Analyzed"))
                    
                volume = birdeye_price.get("volume_24h")
                if volume:
                    metrics_data.append(("Volume 24h", f"${float(volume):,.0f}", "Analyzed"))
            
            # AI risk assessments
            market_metrics = ai_analysis.get("market_metrics", {}) if ai_analysis else {}
            for metric, risk_level in market_metrics.items():
                if metric != "timing_analysis" and risk_level != "unknown":
                    display_name = metric.replace("_", " ").title()
                    metrics_data.append((display_name, "Assessed", risk_level.upper()))
            
            # Add rows
            for metric, value, assessment in metrics_data[:8]:  # Show top 8 metrics
                row = table.add_row()
                row.cells[0].text = str(metric)
                row.cells[1].text = str(value)
                row.cells[2].text = str(assessment)
            
        except Exception as e:
            logger.warning(f"Failed to add discovery scoring table: {e}")
            doc.add_paragraph("Scoring table unavailable")

    def _add_discovery_pros_cons_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add pros/cons section for discovery report"""
        try:
            doc.add_heading('Investment Analysis', level=1)
            
            ai_analysis = analysis_data.get("ai_analysis", {})
            
            # Get pros and cons from AI analysis
            pros = ai_analysis.get("key_insights", [])[:3] if ai_analysis else []
            cons = ai_analysis.get("risk_factors", [])[:3] if ai_analysis else []
            
            # Fallback if no AI data
            if not pros:
                pros = ["Analysis completed", "Security checks performed", "Market data available"]
            if not cons:
                cons = ["Limited historical data", "Market volatility risk", "Requires monitoring"]
            
            # Create table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Headers
            header_row = table.rows[0]
            header_row.cells[0].text = "TOP 3 REASONS TO BUY"
            header_row.cells[1].text = "TOP 3 REASONS TO AVOID"
            
            # Style headers
            header_row.cells[0].paragraphs[0].runs[0].bold = True
            header_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 197, 94)
            header_row.cells[1].paragraphs[0].runs[0].bold = True
            header_row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 68, 68)
            
            # Add content rows
            max_items = max(len(pros), len(cons))
            for i in range(max_items):
                row = table.add_row()
                
                if i < len(pros):
                    row.cells[0].text = f"• {pros[i]}"
                else:
                    row.cells[0].text = ""
                    
                if i < len(cons):
                    row.cells[1].text = f"• {cons[i]}"
                else:
                    row.cells[1].text = ""
            
        except Exception as e:
            logger.warning(f"Failed to add discovery pros/cons: {e}")
            doc.add_paragraph("Investment analysis unavailable")

    def _add_discovery_ai_reasoning_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add AI reasoning section for discovery report"""
        try:
            doc.add_heading('AI Investment Analysis', level=1)
            
            ai_analysis = analysis_data.get("ai_analysis", {})
            
            if ai_analysis:
                reasoning = ai_analysis.get("llama_reasoning", "")
                confidence = ai_analysis.get("confidence", 0)
                
                if reasoning:
                    reasoning_p = doc.add_paragraph()
                    reasoning_p.add_run('AI Analysis: ').bold = True
                    reasoning_p.add_run(reasoning)
                    
                    conf_p = doc.add_paragraph()
                    conf_p.add_run('Analysis Confidence: ').bold = True
                    conf_p.add_run(f'{confidence}%')
                else:
                    doc.add_paragraph("AI reasoning not available")
            else:
                doc.add_paragraph("AI analysis not performed for this token")
            
        except Exception as e:
            logger.warning(f"Failed to add discovery AI reasoning: {e}")
            doc.add_paragraph("AI reasoning unavailable")

    def _add_discovery_market_data_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add market data section for discovery report - FIXED VERSION"""
        try:
            doc.add_heading('Market Data', level=1)
            
            # Use the comprehensive market data extraction (like analysis_storage does)
            market_data = self._extract_comprehensive_market_data(analysis_data)
            
            if not any(market_data.values()):
                doc.add_paragraph("Market data not available from any source")
                return
            
            # Create market data table
            table = doc.add_table(rows=1, cols=3)  # Add source column
            table.style = 'Table Grid'
            
            # Headers
            header_row = table.rows[0]
            header_row.cells[0].text = "Metric"
            header_row.cells[1].text = "Value"
            header_row.cells[2].text = "Source"
            
            # Build market data with sources
            market_data_rows = [
                ("Price", f"${market_data['price_usd']:.8f}" if market_data['price_usd'] else "N/A", market_data.get('price_source', 'N/A')),
                ("Market Cap", f"${market_data['market_cap']:,.0f}" if market_data['market_cap'] else "N/A", market_data.get('market_cap_source', 'N/A')),
                ("24h Volume", f"${market_data['volume_24h']:,.0f}" if market_data['volume_24h'] else "N/A", market_data.get('volume_source', 'N/A')),
                ("Liquidity", f"${market_data['liquidity']:,.0f}" if market_data['liquidity'] else "N/A", market_data.get('liquidity_source', 'N/A')),
                ("24h Change", f"{market_data['price_change_24h']:+.2f}%" if market_data['price_change_24h'] is not None else "N/A", market_data.get('price_change_source', 'N/A'))
            ]
            
            for metric, value, source in market_data_rows:
                row = table.add_row()
                row.cells[0].text = metric
                row.cells[1].text = str(value)
                row.cells[2].text = str(source)
            
            logger.info(f"DOCX market data table created with comprehensive extraction")
            
        except Exception as e:
            logger.warning(f"Failed to add discovery market data: {e}")
            doc.add_paragraph("Market data section unavailable")

    def _extract_comprehensive_market_data(self, analysis_data: Dict[str, Any]) -> Dict[str, Any]:
        """Extract market data with fallback sources (copy of analysis_storage logic)"""
        market_data = {
            "price_usd": None,
            "price_change_24h": None,
            "volume_24h": None,
            "market_cap": None,
            "liquidity": None
        }
        
        try:
            service_responses = analysis_data.get("service_responses", {})
            
            # Primary source: Birdeye
            birdeye_data = service_responses.get("birdeye", {})
            if birdeye_data and birdeye_data.get("price"):
                birdeye_price = birdeye_data["price"]
                
                if birdeye_price.get("value") is not None:
                    market_data["price_usd"] = float(birdeye_price["value"])
                    market_data["price_source"] = "Birdeye"
                
                if birdeye_price.get("price_change_24h") is not None:
                    market_data["price_change_24h"] = float(birdeye_price["price_change_24h"])
                    market_data["price_change_source"] = "Birdeye"
                
                if birdeye_price.get("volume_24h") is not None:
                    market_data["volume_24h"] = float(birdeye_price["volume_24h"])
                    market_data["volume_source"] = "Birdeye"
                
                if birdeye_price.get("market_cap") is not None:
                    market_data["market_cap"] = float(birdeye_price["market_cap"])
                    market_data["market_cap_source"] = "Birdeye"
                
                if birdeye_price.get("liquidity") is not None:
                    market_data["liquidity"] = float(birdeye_price["liquidity"])
                    market_data["liquidity_source"] = "Birdeye"
            
            # Fallback: DexScreener
            if not market_data["volume_24h"] or not market_data["market_cap"]:
                dexscreener_data = service_responses.get("dexscreener", {})
                if dexscreener_data and dexscreener_data.get("pairs", {}).get("pairs"):
                    pairs = dexscreener_data["pairs"]["pairs"]
                    if pairs and len(pairs) > 0:
                        pair = pairs[0]
                        
                        if not market_data["market_cap"] and pair.get("marketCap"):
                            market_data["market_cap"] = float(pair["marketCap"])
                            market_data["market_cap_source"] = "DexScreener"
                        
                        if not market_data["volume_24h"]:
                            vol_data = pair.get("volume", {})
                            if vol_data and vol_data.get("h24"):
                                market_data["volume_24h"] = float(vol_data["h24"])
                                market_data["volume_source"] = "DexScreener"
                        
                        if not market_data["liquidity"]:
                            liq_data = pair.get("liquidity", {})
                            if liq_data and liq_data.get("usd"):
                                market_data["liquidity"] = float(liq_data["usd"])
                                market_data["liquidity_source"] = "DexScreener"
            
            # Fallback: SolSniffer for market cap
            if not market_data["market_cap"]:
                solsniffer_data = service_responses.get("solsniffer", {})
                if solsniffer_data:
                    for field_name in ['marketCap', 'market_cap']:
                        mc_value = solsniffer_data.get(field_name)
                        if mc_value is not None:
                            market_data["market_cap"] = float(mc_value)
                            market_data["market_cap_source"] = "SolSniffer"
                            break
            
            return market_data
            
        except Exception as e:
            logger.warning(f"Error extracting comprehensive market data: {e}")
            return market_data

    def _add_discovery_security_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add security section for discovery report"""
        try:
            doc.add_heading('Security Analysis', level=1)
            
            security_analysis = analysis_data.get("security_analysis", {})
            
            if not security_analysis:
                doc.add_paragraph("Security analysis not available")
                return
            
            # Security status
            overall_safe = security_analysis.get("overall_safe", False)
            critical_issues = security_analysis.get("critical_issues", [])
            warnings = security_analysis.get("warnings", [])
            
            status_p = doc.add_paragraph()
            status_p.add_run('Security Status: ').bold = True
            
            status_run = status_p.add_run("PASSED" if overall_safe else "FAILED")
            status_run.bold = True
            status_run.font.color.rgb = RGBColor(34, 197, 94) if overall_safe else RGBColor(239, 68, 68)
            
            # Security details table
            sec_table = doc.add_table(rows=1, cols=2)
            sec_table.style = 'Table Grid'
            
            # Headers
            sec_header = sec_table.rows[0]
            sec_header.cells[0].text = "Security Check"
            sec_header.cells[1].text = "Result"
            
            security_info = [
                ("Overall Status", "PASSED" if overall_safe else "FAILED"),
                ("Critical Issues", len(critical_issues)),
                ("Warnings", len(warnings)),
                ("Security Score", f"{security_analysis.get('security_score', 'N/A')}/100" if security_analysis.get('security_score') else "N/A")
            ]
            
            for check, result in security_info:
                row = sec_table.add_row()
                row.cells[0].text = str(check)
                row.cells[1].text = str(result)
            
            # List critical issues if any
            if critical_issues:
                doc.add_paragraph("\nCritical Issues:", style='Heading 2')
                for issue in critical_issues[:5]:  # Show top 5
                    issue_p = doc.add_paragraph(f"• {issue}")
                    issue_p.runs[0].font.color.rgb = RGBColor(239, 68, 68)
            
            # List warnings if any
            if warnings:
                doc.add_paragraph("\nWarnings:", style='Heading 2')
                for warning in warnings[:5]:  # Show top 5
                    warning_p = doc.add_paragraph(f"• {warning}")
                    warning_p.runs[0].font.color.rgb = RGBColor(245, 158, 11)
            
        except Exception as e:
            logger.warning(f"Failed to add discovery security section: {e}")
            doc.add_paragraph("Security analysis section unavailable")
    
    def _build_report(self, doc: Document, analysis_data: Dict[str, Any]):
        """Build the enhanced DOCX report structure with AI analysis"""

        try:
            # Extract token info with fallbacks
            token_symbol = self._get_token_symbol(analysis_data)
            token_name = self._get_token_name(analysis_data)
            token_address = analysis_data.get("token_address", "N/A")
            
            # Header
            heading = doc.add_heading(f'Token Analysis: {token_symbol}', 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Token Info
            p = doc.add_paragraph()
            p.add_run('Token Name: ').bold = True
            p.add_run(f'{token_name}\n')
            p.add_run('Symbol: ').bold = True
            p.add_run(f'{token_symbol}\n')
            p.add_run('Contract Address: ').bold = True
            p.add_run(f'{token_address}\n')
            p.add_run('Generated: ').bold = True
            p.add_run(f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
            
            # AI VERDICT SECTION
            self._add_ai_verdict_section(doc, analysis_data)
            
            self._add_enhanced_scoring_table(doc, analysis_data)
            
            self._add_pros_cons_section(doc, analysis_data)
            
            self._add_ai_reasoning_section(doc, analysis_data)
            
            self._add_market_data_safe(doc, analysis_data)
            
            self._add_security_analysis_safe(doc, analysis_data)
            
            self._add_lp_analysis_safe(doc, analysis_data)
            
        except Exception as e:
            logger.error(f"Error building enhanced DOCX report: {e}")
            doc.add_paragraph(f"Error generating full report: {str(e)}")

    def _add_enhanced_scoring_table(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add enhanced scoring table with risk metrics - FIXED"""
        try:
            doc.add_heading('Risk Metrics Analysis', level=1)
            
            # Get enhanced metrics
            overall_analysis = analysis_data.get("overall_analysis", {})
            ai_analysis = analysis_data.get("ai_analysis", {})
            service_responses = analysis_data.get("service_responses", {})
            
            # Create enhanced metrics table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            header_row = table.rows[0]
            header_row.cells[0].text = "Risk Metric"
            header_row.cells[1].text = "Value"
            header_row.cells[2].text = "AI Risk Level"
            
            # Collect metrics from various sources
            metrics_data = []
            
            # Volatility - FIXED extraction
            volatility = None
            if overall_analysis.get("volatility", {}).get("recent_volatility_percent") is not None:
                volatility = overall_analysis["volatility"]["recent_volatility_percent"]
                ai_volatility_risk = ai_analysis.get("market_metrics", {}).get("volatility_risk", "unknown")
                metrics_data.append(("Volatility", f"{volatility}%", ai_volatility_risk.upper()))
            
            # Whale Analysis - FIXED extraction
            whale_data = overall_analysis.get("whale_analysis", {})
            if whale_data and whale_data.get("whale_count") is not None:
                whale_control = whale_data.get("whale_control_percent", 0)
                whale_count = whale_data.get("whale_count", 0)
                ai_whale_risk = ai_analysis.get("market_metrics", {}).get("whale_risk", "unknown")
                metrics_data.append(("Whale Control", f"{whale_count} whales ({whale_control}%)", ai_whale_risk.upper()))
            
            # Volume/Liquidity Ratio - FIXED extraction
            vol_liq_ratio = None
            birdeye = service_responses.get("birdeye", {}).get("price", {})
            if birdeye.get("volume_24h") and birdeye.get("liquidity"):
                try:
                    vol_liq_ratio = (float(birdeye["volume_24h"]) / float(birdeye["liquidity"])) * 100
                    ai_liquidity_health = ai_analysis.get("market_metrics", {}).get("liquidity_health", "unknown")
                    metrics_data.append(("Volume/Liquidity", f"{vol_liq_ratio:.1f}%", ai_liquidity_health.upper()))
                except:
                    pass
            
            # Sniper Risk - FIXED extraction
            sniper_data = overall_analysis.get("sniper_detection", {})
            if sniper_data and sniper_data.get("similar_holders") is not None:
                similar_holders = sniper_data.get("similar_holders", 0)
                pattern_detected = sniper_data.get("pattern_detected", False)
                ai_sniper_risk = ai_analysis.get("market_metrics", {}).get("sniper_risk", "unknown")
                metrics_data.append(("Sniper Patterns", f"{similar_holders} similar ({'Yes' if pattern_detected else 'No'})", ai_sniper_risk.upper()))
            
            # Dev Holdings - FIXED extraction
            dev_percent = self._extract_dev_percent(analysis_data)
            if dev_percent is not None:
                ai_dev_risk = ai_analysis.get("market_metrics", {}).get("dev_risk", "unknown")
                metrics_data.append(("Dev Holdings", f"{dev_percent:.1f}%", ai_dev_risk.upper()))
            
            # LP Security
            lp_status = self._extract_lp_status(analysis_data)
            ai_lp_security = ai_analysis.get("market_metrics", {}).get("lp_security", "unknown")
            metrics_data.append(("LP Security", lp_status, ai_lp_security.upper()))
            
            # Add rows to table
            for metric, value, ai_risk in metrics_data:
                row = table.add_row()
                row.cells[0].text = metric
                row.cells[1].text = str(value)
                row.cells[2].text = str(ai_risk)
                
                # FIXED: Color code AI risk levels properly
                try:
                    risk_cell = row.cells[2]
                    if ai_risk.upper() == "LOW":
                        risk_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 197, 94)  # Green
                    elif ai_risk.upper() == "HIGH":
                        risk_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 68, 68)  # Red
                    elif ai_risk.upper() == "MEDIUM":
                        risk_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)  # Yellow
                except Exception as color_error:
                    logger.debug(f"Color coding failed: {color_error}")
            
            logger.info(f"Enhanced scoring table created with {len(metrics_data)} metrics")
            
        except Exception as e:
            logger.warning(f"Failed to add enhanced scoring table: {e}")
            doc.add_paragraph("Enhanced scoring section unavailable")

    def _add_pros_cons_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add top 3 pros and cons section"""
        try:
            doc.add_heading('Investment Analysis', level=1)
            
            # Get pros and cons from analysis
            pros, cons = self._extract_pros_cons(analysis_data)
            
            # Create pros/cons table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            header_row = table.rows[0]
            header_row.cells[0].text = "TOP 3 REASONS TO BUY"
            header_row.cells[1].text = "TOP 3 REASONS TO AVOID"
            
            # FIXED: Make headers bold and colored properly
            try:
                header_row.cells[0].paragraphs[0].runs[0].bold = True
                header_row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 197, 94)  # Green
                header_row.cells[1].paragraphs[0].runs[0].bold = True  
                header_row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(239, 68, 68)  # Red
            except Exception:
                logger.debug("Header formatting failed")
            
            # Add content rows
            max_items = max(len(pros), len(cons))
            for i in range(max_items):
                row = table.add_row()
                
                # Add pro (if available)
                if i < len(pros):
                    row.cells[0].text = f"• {pros[i]}"
                else:
                    row.cells[0].text = ""
                
                # Add con (if available)
                if i < len(cons):
                    row.cells[1].text = f"• {cons[i]}"
                else:
                    row.cells[1].text = ""
            
            logger.info(f"Pros/cons section added: {len(pros)} pros, {len(cons)} cons")
            
        except Exception as e:
            logger.warning(f"Failed to add pros/cons section: {e}")
            doc.add_paragraph("Pros/cons section unavailable")

    def _add_ai_verdict_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add AI verdict section with GO/WATCH/NO decision - FIXED"""
        try:
            doc.add_heading('AI Investment Verdict', level=1)
            
            # Get verdict from analysis
            overall_analysis = analysis_data.get("overall_analysis", {})
            ai_analysis = analysis_data.get("ai_analysis", {})
            
            # Determine verdict source and decision
            verdict_decision = "UNKNOWN"
            verdict_source = "Traditional Analysis"
            
            # Try to get verdict from AI analysis first (deep analysis)
            if ai_analysis and ai_analysis.get("verdict"):
                verdict_data = ai_analysis["verdict"]
                verdict_decision = verdict_data.get("decision", "UNKNOWN")
                verdict_source = "AI Enhanced Analysis"
            elif overall_analysis and overall_analysis.get("verdict"):
                verdict_data = overall_analysis["verdict"] 
                verdict_decision = verdict_data.get("decision", "UNKNOWN")
            else:
                # Fallback: calculate from score
                score = overall_analysis.get("score", 0)
                if score >= 80:
                    verdict_decision = "GO"
                elif score >= 60:
                    verdict_decision = "WATCH"
                else:
                    verdict_decision = "NO"
            
            # Create verdict display
            p = doc.add_paragraph()
            p.add_run('INVESTMENT VERDICT: ').bold = True
            
            verdict_run = p.add_run(verdict_decision)
            verdict_run.bold = True
            # FIXED: Don't set font.size to RGBColor
            
            # Color coding
            if verdict_decision == "GO":
                verdict_run.font.color.rgb = RGBColor(34, 197, 94)  # Green
            elif verdict_decision == "WATCH":
                verdict_run.font.color.rgb = RGBColor(245, 158, 11)  # Yellow
            else:  # NO
                verdict_run.font.color.rgb = RGBColor(239, 68, 68)  # Red
            
            # Add score
            score = overall_analysis.get("score", 0)
            normalized_score = score / 100  # Convert to 0-1 scale
            
            p.add_run(f'\nAI Score: ').bold = True
            p.add_run(f'{normalized_score:.2f}/1.00 ({score}/100)')
            
            p.add_run(f'\nAnalysis Type: ').bold = True
            p.add_run(f'{verdict_source}')
            
        except Exception as e:
            logger.warning(f"Failed to add AI verdict section: {e}")
            doc.add_paragraph("AI verdict section unavailable")

    def _add_ai_reasoning_section(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add AI-generated investment reasoning section"""
        try:
            doc.add_heading('AI Investment Analysis', level=1)
            
            # Get AI reasoning
            ai_reasoning = self._extract_ai_reasoning(analysis_data)
            
            if ai_reasoning and len(ai_reasoning) > 20:
                # Add AI reasoning text
                reasoning_paragraph = doc.add_paragraph()
                reasoning_paragraph.add_run('AI Analysis: ').bold = True
                reasoning_paragraph.add_run(ai_reasoning)
                
                # Add confidence level
                confidence = self._extract_confidence(analysis_data)
                if confidence:
                    conf_paragraph = doc.add_paragraph()
                    conf_paragraph.add_run('Analysis Confidence: ').bold = True
                    conf_paragraph.add_run(f'{confidence}%')
                
                logger.info("AI reasoning section added successfully")
            else:
                doc.add_paragraph("AI reasoning not available - analysis may have used traditional methods only")
                
        except Exception as e:
            logger.warning(f"Failed to add AI reasoning section: {e}")
            doc.add_paragraph("AI reasoning section unavailable")

    def _extract_pros_cons(self, analysis_data: Dict[str, Any]) -> tuple:
        """Extract top 3 pros and cons from analysis data"""
        try:
            pros = []
            cons = []
            
            # Method 1: Try to get from AI analysis verdict reasoning (deep analysis)
            ai_analysis = analysis_data.get("ai_analysis", {})
            if ai_analysis and ai_analysis.get("verdict", {}).get("reasoning"):
                reasoning = ai_analysis["verdict"]["reasoning"]
                if isinstance(reasoning, dict):
                    pros = reasoning.get("pros", [])[:3]
                    cons = reasoning.get("cons", [])[:3]
                    logger.info("Pros/cons extracted from AI analysis reasoning")
                    return pros, cons
            
            # Method 2: Try to get from overall analysis verdict (quick analysis)
            overall_analysis = analysis_data.get("overall_analysis", {})
            if overall_analysis and overall_analysis.get("verdict", {}).get("reasoning"):
                reasoning = overall_analysis["verdict"]["reasoning"]
                if isinstance(reasoning, dict):
                    pros = reasoning.get("pros", [])[:3]
                    cons = reasoning.get("cons", [])[:3]
                    logger.info("Pros/cons extracted from overall analysis reasoning")
                    return pros, cons
            
            # Method 3: Fallback - construct from available data
            positive_signals = overall_analysis.get("positive_signals", [])
            risk_factors = overall_analysis.get("risk_factors", [])
            
            # Build pros from positive signals and metrics
            if positive_signals:
                pros.extend(positive_signals[:2])
            
            # Add specific metric pros
            if overall_analysis.get("whale_analysis", {}).get("whale_count", 0) == 0:
                pros.append("Perfect token distribution - no whales")
            elif overall_analysis.get("whale_analysis", {}).get("whale_risk_level") == "low":
                pros.append("Low whale concentration risk")
            
            if overall_analysis.get("volatility", {}).get("recent_volatility_percent"):
                vol = overall_analysis["volatility"]["recent_volatility_percent"]
                if vol <= 15:
                    pros.append(f"Low recent volatility ({vol}%)")
            
            # Build cons from risk factors and metrics
            if risk_factors:
                cons.extend(risk_factors[:2])
            
            # Add specific metric cons
            whale_control = overall_analysis.get("whale_analysis", {}).get("whale_control_percent", 0)
            if whale_control > 50:
                cons.append(f"High whale concentration ({whale_control}%)")
            
            if overall_analysis.get("sniper_detection", {}).get("pattern_detected"):
                cons.append("Potential bot/sniper activity detected")
            
            # Ensure we have at least something
            pros = pros[:3] if pros else ["Security checks passed"]
            cons = cons[:3] if cons else ["Limited analysis data"]
            
            logger.info("Pros/cons extracted from fallback method")
            return pros, cons
            
        except Exception as e:
            logger.warning(f"Failed to extract pros/cons: {e}")
            return ["Analysis completed"], ["Limited data available"]
        
    def _extract_ai_reasoning(self, analysis_data: Dict[str, Any]) -> str:
        """Extract AI reasoning from analysis data"""
        try:
            # Try AI analysis first
            ai_analysis = analysis_data.get("ai_analysis", {})
            if ai_analysis:
                # Try llama_reasoning field
                reasoning = ai_analysis.get("llama_reasoning")
                if reasoning and len(reasoning) > 50:  # Meaningful reasoning
                    return reasoning
                
                # Try to construct from AI metrics
                market_metrics = ai_analysis.get("market_metrics", {})
                if market_metrics:
                    risk_assessments = []
                    for metric, risk_level in market_metrics.items():
                        if risk_level and risk_level != "unknown":
                            risk_assessments.append(f"{metric.replace('_', ' ')}: {risk_level}")
                    
                    if risk_assessments:
                        return f"AI risk assessment: {'; '.join(risk_assessments[:5])}"
            
            # Fallback to overall analysis summary
            overall_analysis = analysis_data.get("overall_analysis", {})
            if overall_analysis:
                summary = overall_analysis.get("summary", "")
                if summary and len(summary) > 20:
                    return f"Analysis summary: {summary}"
            
            return "AI analysis completed with available token data"
            
        except Exception as e:
            logger.warning(f"Failed to extract AI reasoning: {e}")
            return "AI reasoning extraction failed"
        
    def _extract_confidence(self, analysis_data: Dict[str, Any]) -> Optional[float]:
        """Extract confidence level from analysis"""
        try:
            # Try AI analysis confidence first
            ai_analysis = analysis_data.get("ai_analysis", {})
            if ai_analysis and ai_analysis.get("confidence"):
                return float(ai_analysis["confidence"])
            
            # Fallback to overall analysis confidence
            overall_analysis = analysis_data.get("overall_analysis", {})
            if overall_analysis and overall_analysis.get("confidence"):
                return float(overall_analysis["confidence"])
            
            return None
            
        except Exception:
            return None

    def _extract_dev_percent(self, analysis_data: Dict[str, Any]) -> Optional[float]:
        """Extract dev holdings percentage"""
        try:
            # Try from service responses
            rugcheck = analysis_data.get("service_responses", {}).get("rugcheck", {})
            if rugcheck and rugcheck.get("creator_analysis"):
                creator_analysis = rugcheck["creator_analysis"]
                creator_balance = creator_analysis.get("creator_balance", 0)
                
                # Try to get total supply
                helius = analysis_data.get("service_responses", {}).get("helius", {})
                if helius and helius.get("supply"):
                    total_supply = helius["supply"].get("ui_amount")
                    if total_supply and creator_balance:
                        return (float(creator_balance) / float(total_supply)) * 100
            
            return None
            
        except Exception:
            return None
        
    def _add_lp_analysis_safe(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add LP analysis with error handling"""
        try:
            doc.add_heading('Liquidity Provider Analysis', level=1)
            
            rugcheck = analysis_data.get("service_responses", {}).get("rugcheck", {})
            
            if rugcheck:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                lp_data = [
                    ("LP Status", self._extract_lp_status(rugcheck)),
                    ("LP Providers", rugcheck.get('total_LP_providers', 'N/A')),
                    ("Locked Value", self._get_locked_value(rugcheck)),
                ]
                
                for key, value in lp_data:
                    row = table.add_row()
                    row.cells[0].text = key
                    row.cells[1].text = str(value)
            else:
                doc.add_paragraph("LP analysis data not available")
        except Exception as e:
            logger.warning(f"Failed to add LP analysis: {e}")
            doc.add_paragraph("LP analysis section unavailable")

    def _extract_lp_status(self, rugcheck_data: Dict[str, Any]) -> str:
        """Determine LP status from RugCheck data"""
        try:
            lockers_data = rugcheck_data.get('lockers_data', {})
            if lockers_data and lockers_data.get('lockers'):
                return "LOCKED"
            
            # Check for burn patterns
            markets = rugcheck_data.get('market_analysis', {}).get('markets', [])
            for market in markets:
                if isinstance(market, dict) and market.get('lp'):
                    holders = market['lp'].get('holders', [])
                    for holder in holders:
                        if isinstance(holder, dict):
                            owner = str(holder.get('owner', ''))
                            pct = holder.get('pct', 0)
                            
                            burn_patterns = ['111111', 'dead', 'burn']
                            if any(pattern in owner.lower() for pattern in burn_patterns) and pct > 50:
                                return "BURNED"
            
            return "UNKNOWN"
            
        except Exception:
            return "UNKNOWN"
        
    
    def _get_locked_value(self, rugcheck_data: Dict[str, Any]) -> str:
        """Get total locked value from RugCheck data"""
        try:
            lockers_data = rugcheck_data.get('lockers_data', {})
            if lockers_data and lockers_data.get('lockers'):
                total_value = 0
                locker_count = 0
                
                for locker_info in lockers_data['lockers'].values():
                    if isinstance(locker_info, dict):
                        usd_locked = locker_info.get('usdcLocked', 0)
                        if isinstance(usd_locked, (int, float)) and usd_locked > 0:
                            total_value += usd_locked
                            locker_count += 1
                
                if total_value > 0:
                    return f"${total_value:,.0f} ({locker_count} lockers)"
                else:
                    return "N/A"
            
            return "N/A"
            
        except Exception as e:
            logger.warning(f"Error extracting locked value: {e}")
            return "N/A"

    def _calculate_rating(self, analysis_data: Dict[str, Any]) -> Dict[str, str]:
        """Calculate GO/WATCH/NO rating based on enhanced data"""
        
        security_data = analysis_data.get("security_analysis", {})
        overall_analysis = analysis_data.get("overall_analysis", {})
        
        # Critical security failure = NO
        if not security_data.get("overall_safe"):
            return {
                "decision": "NO",
                "reasoning": "Critical security issues detected"
            }
        
        # Try to get verdict from analysis
        verdict = None
        if overall_analysis.get("verdict"):
            verdict = overall_analysis["verdict"].get("decision")
        
        # If no verdict, calculate from score
        if not verdict:
            score = overall_analysis.get("score", 0)
            risk_level = overall_analysis.get("risk_level", "unknown")
            
            if score >= 80 and risk_level == "low":
                verdict = "GO"
            elif score >= 60:
                verdict = "WATCH"
            else:
                verdict = "NO"
        
        # Generate reasoning
        reasoning_parts = []
        
        # Add score component
        score = overall_analysis.get("score", 0)
        normalized_score = score / 100
        reasoning_parts.append(f"AI Score: {normalized_score:.2f}/1.00")
        
        # Add key factors
        whale_count = overall_analysis.get("whale_analysis", {}).get("whale_count", 0)
        if whale_count == 0:
            reasoning_parts.append("Perfect token distribution")
        elif overall_analysis.get("whale_analysis", {}).get("whale_control_percent", 0) > 50:
            reasoning_parts.append("High whale concentration risk")
        
        volatility = overall_analysis.get("volatility", {}).get("recent_volatility_percent")
        if volatility is not None:
            if volatility <= 10:
                reasoning_parts.append("Low volatility")
            elif volatility > 30:
                reasoning_parts.append("High volatility")
        
        reasoning = " | ".join(reasoning_parts[:3]) if reasoning_parts else f"Score-based assessment ({normalized_score:.2f}/1.00)"
        
        return {
            "decision": verdict,
            "reasoning": reasoning
        }

    def _add_final_rating_safe(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add final rating with enhanced verdict system"""
        try:
            doc.add_heading('Final Investment Rating', level=1)
            
            rating = self._calculate_rating(analysis_data)
            
            p = doc.add_paragraph()
            p.add_run('FINAL RATING: ').bold = True
            
            rating_run = p.add_run(rating["decision"])
            rating_run.bold = True
            if rating["decision"] == "GO":
                rating_run.font.color.rgb = RGBColor(34, 197, 94)  # Green
            elif rating["decision"] == "NO":
                rating_run.font.color.rgb = RGBColor(239, 68, 68)  # Red
            else:  # WATCH
                rating_run.font.color.rgb = RGBColor(245, 158, 11)  # Yellow
            
            # Add reasoning
            doc.add_paragraph(f"Reasoning: {rating['reasoning']}")
            
            # Add disclaimer
            disclaimer = doc.add_paragraph()
            disclaimer.add_run('Disclaimer: ').bold = True
            disclaimer.add_run('This analysis is for informational purposes only and does not constitute financial advice. Cryptocurrency investments carry significant risk.')
            
        except Exception as e:
            logger.warning(f"Failed to add final rating: {e}")
            doc.add_paragraph("Final rating section unavailable")

    def _add_security_analysis_safe(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add security analysis with error handling"""
        try:
            doc.add_heading('Security Analysis', level=1)
            security_data = analysis_data.get("security_analysis", {})
            
            if security_data:
                sec_table = doc.add_table(rows=1, cols=2)
                sec_table.style = 'Table Grid'
                
                security_info = [
                    ("Security Status", "PASSED" if security_data.get("overall_safe") else "FAILED"),
                    ("Critical Issues", len(security_data.get("critical_issues", []))),
                    ("Warnings", len(security_data.get("warnings", []))),
                ]
                
                for key, value in security_info:
                    row = sec_table.add_row()
                    row.cells[0].text = key
                    row.cells[1].text = str(value)
            else:
                doc.add_paragraph("Security analysis data not available")
        except Exception as e:
            logger.warning(f"Failed to add security analysis: {e}")
            doc.add_paragraph("Security analysis section unavailable")

    def _add_market_data_safe(self, doc: Document, analysis_data: Dict[str, Any]):
        """Add market data with error handling - FIXED to use aggregated data"""
        try:
            doc.add_heading('Market Data', level=1)
            
            # Try to get aggregated market data from AI analysis first
            ai_analysis = analysis_data.get("ai_analysis", {})
            market_metrics = ai_analysis.get("market_metrics", {}) if ai_analysis else {}
            
            # Also check if comprehensive market data was calculated in Groq service
            service_responses = analysis_data.get("service_responses", {})
            
            # Initialize market data dict
            market_data = {}

            # Check Birdeye first
            birdeye = service_responses.get("birdeye", {}).get("price", {})
            if birdeye:
                # Price from Birdeye
                if not market_data.get("price") and birdeye.get("value"):
                    try:
                        market_data["price"] = float(birdeye["value"])
                        market_data["price_source"] = "Birdeye"
                    except (ValueError, TypeError):
                        pass
                
                # Market cap from Birdeye
                if not market_data.get("market_cap") and birdeye.get("market_cap"):
                    try:
                        market_data["market_cap"] = float(birdeye["market_cap"])
                        market_data["market_cap_source"] = "Birdeye"
                    except (ValueError, TypeError):
                        pass
                
                # Volume from Birdeye
                if not market_data.get("volume_24h") and birdeye.get("volume_24h"):
                    try:
                        market_data["volume_24h"] = float(birdeye["volume_24h"])
                        market_data["volume_source"] = "Birdeye"
                    except (ValueError, TypeError):
                        pass
                
                # Liquidity from Birdeye
                if not market_data.get("liquidity") and birdeye.get("liquidity"):
                    try:
                        market_data["liquidity"] = float(birdeye["liquidity"])
                        market_data["liquidity_source"] = "Birdeye"
                    except (ValueError, TypeError):
                        pass
                
                # Price change from Birdeye
                if not market_data.get("price_change_24h") and birdeye.get("price_change_24h") is not None:
                    try:
                        market_data["price_change_24h"] = float(birdeye["price_change_24h"])
                        market_data["price_change_source"] = "Birdeye"
                    except (ValueError, TypeError):
                        pass

            # SolSniffer fallback
            if not market_data.get("market_cap"):
                solsniffer = service_responses.get("solsniffer", {})
                if solsniffer and isinstance(solsniffer, dict):
                    for field_name in ['marketCap', 'market_cap']:
                        mc_value = solsniffer.get(field_name)
                        if mc_value is not None:
                            try:
                                market_data["market_cap"] = float(mc_value)
                                market_data["market_cap_source"] = "SolSniffer"
                                break
                            except (ValueError, TypeError):
                                continue
            
            # DexScreener fallback
            dexscreener = service_responses.get("dexscreener", {})
            if dexscreener and dexscreener.get("pairs", {}).get("pairs"):
                pairs = dexscreener["pairs"]["pairs"]
                if isinstance(pairs, list) and len(pairs) > 0:
                    pair = pairs[0]
                    
                    # Market cap from DexScreener
                    if not market_data.get("market_cap") and pair.get("marketCap"):
                        try:
                            market_data["market_cap"] = float(pair["marketCap"])
                            market_data["market_cap_source"] = "DexScreener"
                        except (ValueError, TypeError):
                            pass
                    
                    # Volume from DexScreener
                    if not market_data.get("volume_24h"):
                        vol_data = pair.get("volume", {})
                        if vol_data and vol_data.get("h24"):
                            try:
                                market_data["volume_24h"] = float(vol_data["h24"])
                                market_data["volume_source"] = "DexScreener"
                            except (ValueError, TypeError):
                                pass
                    
                    # Liquidity from DexScreener
                    if not market_data.get("liquidity"):
                        liq_data = pair.get("liquidity", {})
                        if liq_data and liq_data.get("usd"):
                            try:
                                market_data["liquidity"] = float(liq_data["usd"])
                                market_data["liquidity_source"] = "DexScreener"
                            except (ValueError, TypeError):
                                pass
                    
                    # Price from DexScreener
                    if not market_data.get("price") and pair.get("priceUsd"):
                        try:
                            market_data["price"] = float(pair["priceUsd"])
                            market_data["price_source"] = "DexScreener"
                        except (ValueError, TypeError):
                            pass
            
            # Create table with extracted data
            if market_data:
                table = doc.add_table(rows=1, cols=3)  # Add source column
                table.style = 'Table Grid'
                
                # Header row
                header_row = table.rows[0]
                header_row.cells[0].text = "Metric"
                header_row.cells[1].text = "Value"
                header_row.cells[2].text = "Source"
                
                # Market data rows
                market_data_rows = [
                    ("Price", f"${market_data.get('price', 'N/A')}" if market_data.get('price') else "N/A", market_data.get('price_source', 'N/A')),
                    ("Market Cap", f"${market_data.get('market_cap', 0):,.0f}" if market_data.get('market_cap') else "N/A", market_data.get('market_cap_source', 'N/A')),
                    ("24h Volume", f"${market_data.get('volume_24h', 0):,.0f}" if market_data.get('volume_24h') else "N/A", market_data.get('volume_source', 'N/A')),
                    ("Liquidity", f"${market_data.get('liquidity', 0):,.0f}" if market_data.get('liquidity') else "N/A", market_data.get('liquidity_source', 'N/A')),
                    ("24h Change", f"{market_data.get('price_change_24h', 'N/A')}%" if market_data.get('price_change_24h') is not None else "N/A", market_data.get('price_change_source', 'N/A')),
                ]
                
                for metric, value, source in market_data_rows:
                    row = table.add_row()
                    row.cells[0].text = metric
                    row.cells[1].text = str(value)
                    row.cells[2].text = str(source)
                
                logger.info(f"DOCX market data table created with {len(market_data_rows)} rows")
            else:
                doc.add_paragraph("Market data not available from any source")
                logger.warning("No market data available for DOCX generation")
                
        except Exception as e:
            logger.warning(f"Failed to add market data: {e}")
            doc.add_paragraph("Market data section unavailable")
    
    def _get_token_symbol(self, analysis_data: Dict[str, Any]) -> str:
        """Extract token symbol from various sources - IMPROVED with fallbacks"""
        try:
            services = analysis_data.get("service_responses", {})
            
            # Try SolSniffer first
            if services.get("solsniffer", {}).get("tokenSymbol"):
                return services["solsniffer"]["tokenSymbol"]
            
            # Try Helius
            helius_meta = services.get("helius", {}).get("metadata", {})
            if helius_meta.get("onChainMetadata", {}).get("metadata", {}).get("data", {}).get("symbol"):
                return helius_meta["onChainMetadata"]["metadata"]["data"]["symbol"]
            
            # Try other sources
            for service_name, service_data in services.items():
                if isinstance(service_data, dict):
                    for key in ["symbol", "tokenSymbol", "token_symbol"]:
                        if service_data.get(key):
                            return service_data[key]
            
            return "Unknown"
        except Exception as e:
            logger.warning(f"Error extracting token symbol: {e}")
            return "Unknown"
        
    def _get_token_name(self, analysis_data: Dict[str, Any]) -> str:
        """Extract token name from various sources - IMPROVED with fallbacks"""
        try:
            services = analysis_data.get("service_responses", {})
            
            # Try SolSniffer first
            if services.get("solsniffer", {}).get("tokenName"):
                return services["solsniffer"]["tokenName"]
            
            # Try Helius
            helius_meta = services.get("helius", {}).get("metadata", {})
            if helius_meta.get("onChainMetadata", {}).get("metadata", {}).get("data", {}).get("name"):
                return helius_meta["onChainMetadata"]["metadata"]["data"]["name"]
            
            # Try other sources
            for service_name, service_data in services.items():
                if isinstance(service_data, dict):
                    for key in ["name", "tokenName", "token_name"]:
                        if service_data.get(key):
                            return service_data[key]
            
            return "Unknown Token"
        except Exception as e:
            logger.warning(f"Error extracting token name: {e}")
            return "Unknown Token"

# Global service instance
docx_service = DocxReportService()